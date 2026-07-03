from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT = Path(__file__).with_name("oferta_send_invite.docx")

FONT = "Calibri"
INK = RGBColor(36, 33, 35)
MUTED = RGBColor(104, 96, 101)
ACCENT = RGBColor(184, 49, 81)
LIGHT_FILL = "F7F3F5"
TABLE_BORDER = "D9D1D6"


def set_run_font(
    run,
    *,
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: RGBColor | None = None,
):
    run.font.name = FONT
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_margins(cell, *, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Страница ")
    set_run_font(run, size=9, color=MUTED)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_multilevel_numbering(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)

    level_specs = [
        (0, "%1.", 0, 0),
        (1, "%1.%2.", 0, 360),
    ]
    for ilvl, text, left, hanging in level_specs:
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(ilvl))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "decimal")
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        lvl.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "space")
        lvl.append(suff)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left + hanging))
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left + hanging))
        ind.set(qn("w:hanging"), str(hanging))
        p_pr.append(ind)
        lvl.append(p_pr)
        abstract.append(lvl)

    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int, level: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)


def add_section_heading(doc: Document, num_id: int, text: str):
    p = doc.add_paragraph(style="Heading 1")
    apply_numbering(p, num_id, 0)
    run = p.add_run(text)
    set_run_font(run, size=14, bold=True, color=ACCENT)
    return p


def add_clause(doc: Document, num_id: int, text: str):
    p = doc.add_paragraph(style="Clause")
    apply_numbering(p, num_id, 1)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=INK)
    return p


def add_requisites_table(doc: Document):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    rows = [
        ("Статус", "Плательщик налога на профессиональный доход (самозанятый)"),
        ("ФИО Исполнителя", "[УКАЖИТЕ ПОЛНЫЕ ФИО ИСПОЛНИТЕЛЯ]"),
        ("ИНН", "780527120543"),
        ("Сайт", "https://send-invite.online"),
        ("Телефон", "+7 904 217-63-94"),
        ("E-mail", "Synkoveugeny@yandex.ru"),
    ]
    for label, value in rows:
        cells = table.add_row().cells
        left = cells[0].paragraphs[0]
        left.paragraph_format.space_after = Pt(0)
        left_run = left.add_run(label)
        set_run_font(left_run, size=10, bold=True, color=INK)
        right = cells[1].paragraphs[0]
        right.paragraph_format.space_after = Pt(0)
        right_run = right.add_run(value)
        set_run_font(
            right_run,
            size=10,
            bold=label == "ФИО Исполнителя",
            color=ACCENT if label == "ФИО Исполнителя" else INK,
        )
        if label == "ФИО Исполнителя":
            right_run.font.highlight_color = 7
        set_cell_shading(cells[0], LIGHT_FILL)
    set_table_geometry(table, [2700, 6660])
    return table


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.885)
    section.right_margin = Inches(0.885)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    heading = styles["Heading 1"]
    heading.font.name = FONT
    heading._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    heading._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    heading.font.size = Pt(14)
    heading.font.bold = True
    heading.font.color.rgb = ACCENT
    heading.paragraph_format.space_before = Pt(14)
    heading.paragraph_format.space_after = Pt(7)
    heading.paragraph_format.keep_with_next = True

    clause_style = styles.add_style("Clause", WD_STYLE_TYPE.PARAGRAPH)
    clause_style.font.name = FONT
    clause_style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    clause_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    clause_style.font.size = Pt(10.5)
    clause_style.font.color.rgb = INK
    clause_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    clause_style.paragraph_format.space_after = Pt(5)
    clause_style.paragraph_format.line_spacing = 1.1
    clause_style.paragraph_format.widow_control = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("SEND INVITE  |  ПУБЛИЧНАЯ ОФЕРТА")
    set_run_font(run, size=8.5, bold=True, color=MUTED)

    add_page_number(section.footer.paragraphs[0])

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("ПУБЛИЧНАЯ ОФЕРТА")
    set_run_font(title_run, size=22, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(6)
    subtitle_run = subtitle.add_run("о заключении договора возмездного оказания услуг")
    set_run_font(subtitle_run, size=12, color=MUTED)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(18)
    meta_run = meta.add_run("Редакция от 02 июля 2026 года  •  https://send-invite.online")
    set_run_font(meta_run, size=9.5, color=ACCENT, bold=True)

    lead = doc.add_paragraph()
    lead.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    lead.paragraph_format.space_after = Pt(10)
    lead_run = lead.add_run(
        "Настоящий документ является официальным предложением Исполнителя заключить договор "
        "на оказание услуг по созданию и публикации персонального сайта-приглашения. "
        "До оплаты Заказчику рекомендуется внимательно ознакомиться с условиями Оферты, "
        "описанием выбранного тарифа и Политикой обработки персональных данных."
    )
    set_run_font(lead_run, size=10.5, color=INK)

    num_id = add_multilevel_numbering(doc)

    sections: list[tuple[str, list[str]]] = [
        (
            "Общие положения",
            [
                "Настоящая Оферта регулирует отношения между Исполнителем и дееспособным физическим лицом, приобретающим Услугу для личных, семейных и иных нужд, не связанных с предпринимательской деятельностью (далее — «Заказчик»).",
                "Оферта является публичным предложением в соответствии с пунктом 2 статьи 437 Гражданского кодекса Российской Федерации.",
                "Акцептом Оферты является одновременное совершение Заказчиком следующих действий: ознакомление с условиями Оферты и описанием Услуги, подтверждение согласия с ними в интерфейсе Сервиса и оплата выбранного тарифа.",
                "Договор считается заключенным с момента подтверждения оплаты платежным сервисом. Условия Оферты, действовавшие на момент оплаты, применяются к соответствующему Заказу в течение всего срока его исполнения.",
            ],
        ),
        (
            "Термины и определения",
            [
                "Сервис — программный комплекс Send Invite, доступный по адресу https://send-invite.online, включая редактор, личный кабинет и опубликованные сайты-приглашения.",
                "Сайт-приглашение — персональная веб-страница, созданная Заказчиком в Сервисе на основе выбранного шаблона и предоставленных им материалов.",
                "Услуга — предоставление доступа к функциональности Сервиса, обработка введенных Заказчиком данных и публикация Сайта-приглашения по уникальной ссылке на условиях выбранного тарифа.",
                "Тариф — совокупность состава Услуги, стоимости, срока доступности опубликованного Сайта-приглашения и иных условий, отображаемых на Сайте до оплаты.",
                "Заказ — оформленный Заказчиком запрос на оказание Услуги по выбранному тарифу.",
                "Личный кабинет — закрытый раздел Сервиса, доступный Заказчику после авторизации через Yandex ID.",
            ],
        ),
        (
            "Предмет Договора",
            [
                "Исполнитель обязуется оказать Заказчику Услугу по созданию и публикации Сайта-приглашения, а Заказчик обязуется принять и оплатить Услугу.",
                "Конкретный шаблон, состав функций, стоимость, срок доступности, ограничения и иные характеристики Услуги определяются описанием выбранного тарифа, показанным Заказчику до оплаты, и являются неотъемлемой частью Договора.",
                "Услуга не включает разработку индивидуального дизайна, подготовку текстов, обработку фотографий, создание музыки, организацию мероприятия или рассылку приглашений, если иное прямо не указано в описании тарифа.",
                "Исполнитель вправе привлекать третьих лиц для технического оказания Услуги, оставаясь ответственным перед Заказчиком за их действия в пределах, установленных законом.",
            ],
        ),
        (
            "Оформление Заказа",
            [
                "Для оформления Заказа Заказчик выбирает шаблон, авторизуется в Сервисе, заполняет сведения о мероприятии, загружает материалы и проверяет предварительный вид Сайта-приглашения.",
                "Заказчик обязан указать достоверные контактные данные и самостоятельно проверить тексты, даты, адреса, ссылки, фотографии и иные материалы до публикации.",
                "До оплаты Сервис показывает Заказчику наименование Услуги, итоговую стоимость, срок доступности и существенные ограничения выбранного тарифа.",
                "Исполнитель вправе отказать в публикации материалов, нарушающих закон, права третьих лиц, общепринятые нормы или технические требования Сервиса. В таком случае вопрос о возврате оплаты решается с учетом фактически оказанной части Услуги и понесенных расходов.",
            ],
        ),
        (
            "Стоимость и порядок оплаты",
            [
                "Стоимость Услуги указывается в рублях Российской Федерации на Сайте и в интерфейсе оформления Заказа до момента оплаты. Итоговая стоимость фиксируется для соответствующего Заказа после его оплаты.",
                "Оплата производится в безналичном порядке через платежный сервис Robokassa. Доступные способы оплаты определяются платежным сервисом и отображаются на платежной странице.",
                "Обязательство Заказчика по оплате считается исполненным после получения Исполнителем подтверждения об успешном платеже.",
                "Исполнитель применяет специальный налоговый режим «Налог на профессиональный доход». НДС к стоимости Услуги не предъявляется. Электронный чек направляется по контактным данным, указанным Заказчиком при оплате.",
                "Исполнитель не получает и не хранит полные реквизиты банковской карты Заказчика. Обработка платежных данных выполняется платежным сервисом и участвующими кредитными организациями.",
            ],
        ),
        (
            "Порядок и сроки оказания Услуги",
            [
                "Оказание Услуги начинается автоматически после подтверждения оплаты, если иное не указано в описании тарифа.",
                "Результатом Услуги является активация публичной ссылки на Сайт-приглашение и отображение опубликованного проекта в Личном кабинете Заказчика.",
                "Услуга считается оказанной в части создания и публикации с момента активации публичной ссылки. Поддержание доступа к Сайту-приглашению осуществляется в течение срока, указанного в выбранном тарифе.",
                "Срок активации может быть увеличен на период технических работ, сбоев сторонних сервисов, проверки материалов или устранения обстоятельств, зависящих от Заказчика.",
                "Заказчик обязан сообщить о выявленном недостатке по контактному e-mail Исполнителя. Исполнитель устраняет воспроизводимые технические недостатки в разумный срок либо предлагает иной предусмотренный законом способ урегулирования.",
            ],
        ),
        (
            "Права и обязанности Сторон",
            [
                "Исполнитель обязуется обеспечивать доступность оплаченной функциональности в пределах выбранного тарифа, защищать обрабатываемые данные и предоставлять Заказчику достоверную информацию об Услуге.",
                "Исполнитель вправе проводить профилактические и аварийные работы, обновлять Сервис, заменять технические решения и ограничивать доступ при выявлении угроз безопасности или нарушений настоящей Оферты.",
                "Заказчик обязан использовать Сервис добросовестно, не вмешиваться в его работу, не обходить технические ограничения и не передавать доступ к Личному кабинету посторонним лицам.",
                "Заказчик обязан сохранять контроль над своей учетной записью Yandex ID. Действия, совершенные после успешной авторизации, считаются действиями Заказчика, пока Исполнитель не уведомлен о несанкционированном доступе.",
                "Заказчик вправе пользоваться опубликованным Сайтом-приглашением и направлять ссылку гостям в течение оплаченного срока доступа.",
            ],
        ),
        (
            "Материалы Заказчика и интеллектуальные права",
            [
                "Заказчик сохраняет права на загруженные им тексты, изображения, аудиозаписи и иные материалы и предоставляет Исполнителю безвозмездное право использовать их исключительно для исполнения Договора: хранения, технической обработки, отображения и передачи посетителям Сайта-приглашения.",
                "Заказчик гарантирует наличие необходимых прав и согласий на использование загружаемых материалов, включая изображения людей, музыкальные произведения, товарные знаки и персональные данные.",
                "Исключительные права на программный код Сервиса, шаблоны, элементы интерфейса, графику и иные материалы Исполнителя принадлежат Исполнителю или соответствующим правообладателям.",
                "Оплата Услуги не означает передачу Заказчику исключительных прав на Сервис или шаблоны. Заказчику предоставляется ограниченное право использовать результат Услуги в личных некоммерческих целях в течение срока выбранного тарифа.",
            ],
        ),
        (
            "Отказ от Услуги и возврат денежных средств",
            [
                "Заказчик вправе отказаться от исполнения Договора в любое время при условии оплаты Исполнителю фактически понесенных расходов и фактически оказанной части Услуги в соответствии с законодательством Российской Федерации.",
                "До активации публичной ссылки Заказчик вправе направить заявление об отказе и возврате на e-mail Исполнителя. После активации ссылка считается результатом оказанной части Услуги; размер возврата определяется с учетом уже оказанной части и фактических расходов.",
                "При существенном недостатке Услуги, нарушении срока или невозможности использования оплаченной функциональности по вине Исполнителя Заказчик вправе заявить требования, предусмотренные законодательством о защите прав потребителей.",
                "Заявление должно содержать ФИО Заказчика, контактный e-mail, дату и сумму платежа, идентификатор Заказа, причину обращения и требование. Исполнитель вправе запросить сведения, необходимые для идентификации платежа.",
                "Решение по заявлению и возврат денежных средств осуществляются в сроки, установленные законодательством Российской Федерации. Возврат производится, как правило, тем же способом, которым была произведена оплата; срок зачисления зависит от банка и платежной системы.",
            ],
        ),
        (
            "Персональные данные",
            [
                "Обработка персональных данных Заказчика осуществляется в соответствии с Политикой обработки персональных данных, размещенной по адресу https://send-invite.online/privacy.",
                "Для работы Сервиса могут обрабатываться данные учетной записи Yandex ID, контактные данные, сведения о Заказах, платежах, мероприятии, а также техническая информация об использовании Сервиса.",
                "При включении формы RSVP Сервис может обрабатывать данные гостей, введенные ими на Сайте-приглашении. Заказчик обязан обеспечить законность размещения сведений о третьих лицах и информировать приглашенных о порядке обработки их данных.",
                "Для исполнения Договора данные могут передаваться поставщикам авторизации, платежному сервису, хостинг-провайдерам, операторам хранения файлов и другим обработчикам в объеме, необходимом для работы соответствующей функции.",
            ],
        ),
        (
            "Ответственность и ограничения",
            [
                "Стороны несут ответственность за неисполнение обязательств в соответствии с Договором и законодательством Российской Федерации.",
                "Исполнитель не отвечает за ошибки в материалах Заказчика, действия посетителей Сайта-приглашения, сбои оборудования или связи Заказчика, работу Yandex ID, платежного сервиса, банков, операторов связи и иных независимых третьих лиц.",
                "Исполнитель не гарантирует бесперебойную работу Сервиса во все моменты времени, но принимает разумные меры для восстановления доступности и сохранности данных.",
                "Ограничения ответственности, предусмотренные Офертой, не применяются в случаях, когда такое ограничение запрещено императивными нормами законодательства о защите прав потребителей.",
            ],
        ),
        (
            "Обстоятельства непреодолимой силы",
            [
                "Сторона освобождается от ответственности за нарушение обязательства, вызванное чрезвычайными и непредотвратимыми при данных условиях обстоятельствами, включая стихийные бедствия, военные действия, запреты органов власти, масштабные сбои инфраструктуры связи и иные обстоятельства непреодолимой силы.",
                "Сторона, столкнувшаяся с такими обстоятельствами, уведомляет другую Сторону при первой разумной возможности. Срок исполнения обязательств продлевается на период действия обстоятельств и устранения их последствий.",
            ],
        ),
        (
            "Срок действия и изменение Оферты",
            [
                "Оферта вступает в силу с момента публикации на Сайте и действует до ее отзыва Исполнителем.",
                "Исполнитель вправе изменять Оферту путем размещения новой редакции на Сайте. Новая редакция применяется к Заказам, оплаченным после даты ее публикации, если иное прямо не предусмотрено законом.",
                "Договор с конкретным Заказчиком действует с момента Акцепта до полного исполнения обязательств Сторонами, включая окончание оплаченного срока доступа.",
            ],
        ),
        (
            "Разрешение споров и контакты",
            [
                "К отношениям Сторон применяется законодательство Российской Федерации, включая нормы о защите прав потребителей.",
                "Претензии и обращения направляются на e-mail Исполнителя. Исполнитель рассматривает обращение и направляет ответ в срок, установленный законодательством Российской Федерации.",
                "Если спор не урегулирован путем переговоров, он подлежит рассмотрению судом в порядке и по правилам подсудности, установленным законодательством Российской Федерации.",
                "Юридически значимые сообщения могут направляться по адресам электронной почты, указанным Заказчиком и Исполнителем, если закон не требует иной формы.",
            ],
        ),
    ]

    for heading_text, clauses in sections:
        add_section_heading(doc, num_id, heading_text)
        for clause in clauses:
            add_clause(doc, num_id, clause)

    add_section_heading(doc, num_id, "Реквизиты Исполнителя")
    add_requisites_table(doc)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(10)
    note.paragraph_format.space_after = Pt(0)
    note_run = note.add_run(
        "Перед публикацией заполните выделенное поле «ФИО Исполнителя» полными ФИО "
        "самозанятого. Остальные реквизиты перенесены из исходного документа."
    )
    set_run_font(note_run, size=9, italic=True, color=ACCENT)

    doc.core_properties.title = "Публичная оферта Send Invite"
    doc.core_properties.subject = "Договор оказания услуг по созданию и публикации сайта-приглашения"
    doc.core_properties.author = "Send Invite"
    doc.core_properties.keywords = "публичная оферта, услуги, сайт-приглашение, Send Invite"
    doc.core_properties.comments = "Адаптировано для сервиса Send Invite"

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
